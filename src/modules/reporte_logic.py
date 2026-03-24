# src/modules/reporte_logic.py

# Importa utilidades del sistema de archivos
import os

# Importa datetime para manejo de fechas y nombres de archivos
from datetime import datetime

# Importa Tuple para anotaciones de tipo
from typing import Tuple

# Importa clases y utilidades de python-docx para manipular documentos Word
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

# Importa PyMySQL para consultas a la base de datos
import pymysql

# Importa funciones de conexión a base de datos
from modules.conexion import crear_conexion, cerrar_conexion


# Directorio base del proyecto
BASE_DIR = os.path.dirname(os.path.dirname(__file__))

# Ruta completa hacia la plantilla DOCX del reporte
TEMPLATE_PATH = os.path.join(BASE_DIR, "reports", "formato_asis_insur.docx")



def _grado_nombre_corto(grado_seccion: str) -> Tuple[str, str]:
    # Intenta separar el grado y la sección a partir del texto tipo "6-1"
    try:
        parts = grado_seccion.split("-")
        grado_num = int(parts[0])
        seccion = grado_seccion
    except Exception:
        # Si falla el formato esperado, retorna el texto tal como llegó
        return (grado_seccion, grado_seccion)


    # Mapeo de número de grado a nombre corto legible
    mapping = {
        6: "Sexto",
        7: "Séptimo",
        8: "Octavo",
        9: "Noveno",
        10: "Décimo",
        11: "Once"
    }

    # Retorna nombre del grado y la sección
    return (mapping.get(grado_num, f"Grado {grado_num}"), seccion)



def _replace_text_in_paragraphs(doc: Document, marker: str, text: str):
    # Recorre todos los párrafos del documento
    for p in doc.paragraphs:
        if marker in p.text:
            # Guarda los runs actuales del párrafo
            inline = p.runs

            # Reemplaza el marcador por el texto real
            new_text = p.text.replace(marker, text)

            # Elimina los runs existentes para reconstruir el párrafo
            for i in range(len(inline)-1, -1, -1):
                p._element.remove(inline[i]._element)

            # Agrega el nuevo texto ya reemplazado
            p.add_run(new_text)



def _replace_in_tables(doc: Document, marker: str, text: str):
    # Recorre todas las tablas del documento
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if marker in cell.text:
                    # Recorre los párrafos de cada celda para reemplazar el marcador
                    for p in cell.paragraphs:
                        if marker in p.text:
                            new_text = p.text.replace(marker, text)

                            # Elimina los runs existentes de ese párrafo
                            for run in p.runs:
                                p._element.remove(run._element)

                            # Inserta el nuevo contenido
                            p.add_run(new_text)



def generar_reporte_pdf(desde_str: str, hasta_str: str, grado_seccion: str, docente: dict = None, codigo_doc: str = None) -> dict:
    # Convierte las fechas recibidas en texto al formato date
    try:
        desde = datetime.strptime(desde_str, "%d/%m/%Y").date()
        hasta = datetime.strptime(hasta_str, "%d/%m/%Y").date()
    except Exception as e:
        return {"ok": False, "msg": f"Formato de fecha inválido: {e}", "pdf_path": None, "docx_path": None}


    # Abre conexión con la base de datos
    conexion = crear_conexion()
    if not conexion:
        return {"ok": False, "msg": "No se pudo conectar a la base de datos", "pdf_path": None, "docx_path": None}


    # Crea cursor tipo diccionario
    cursor = conexion.cursor(pymysql.cursors.DictCursor)
    try:
        # Obtener alumnos del grado
        cursor.execute(
            """
            SELECT m.id_matricula, e.id_estudiante, e.nombres, e.apellidos
            FROM matriculas m
            INNER JOIN estudiantes e ON m.id_estudiante = e.id_estudiante
            WHERE m.grado = %s AND m.estado = 'Estudiante'
            ORDER BY e.apellidos ASC
            """,
            (grado_seccion,)
        )
        filas = cursor.fetchall()
        if not filas:
            alumnos = []
            id_matriculas = []
        else:
            # Construye una lista normalizada de alumnos
            alumnos = [{
                "id_matricula": f["id_matricula"],
                "id_estudiante": f["id_estudiante"],
                "nombres": f["nombres"] or "",
                "apellidos": f["apellidos"] or ""
            } for f in filas]
            id_matriculas = [a["id_matricula"] for a in alumnos]


        # Obtener fechas de asistencia
        fechas = []
        if id_matriculas:
            # Genera placeholders dinámicos para el IN de matrícula
            format_ids = ",".join(["%s"] * len(id_matriculas))
            q = f"""
                SELECT DISTINCT fecha
                FROM asistencias
                WHERE id_matricula IN ({format_ids})
                  AND fecha BETWEEN %s AND %s
                ORDER BY fecha ASC
            """
            params = id_matriculas + [desde, hasta]
            cursor.execute(q, params)
            fechas_rows = cursor.fetchall()
            fechas = [r["fecha"] for r in fechas_rows]


        # Limita el reporte a máximo 7 fechas
        fechas = fechas[:7]


        # Obtener asistencias
        asist_map = {}
        if id_matriculas and fechas:
            # Genera placeholders dinámicos para matrículas y fechas
            format_ids = ",".join(["%s"] * len(id_matriculas))
            format_fechas = ",".join(["%s"] * len(fechas))
            q2 = f"""
                SELECT id_matricula, fecha, estado
                FROM asistencias
                WHERE id_matricula IN ({format_ids})
                  AND fecha IN ({format_fechas})
            """
            params2 = id_matriculas + fechas
            cursor.execute(q2, params2)
            for r in cursor.fetchall():
                id_mat = r["id_matricula"]
                fecha = r["fecha"]
                estado = r["estado"]
                asist_map[(id_mat, fecha)] = estado
    finally:
        # Cierra cursor y conexión
        cursor.close()
        cerrar_conexion(conexion)


    # Verifica que exista la plantilla del reporte
    if not os.path.exists(TEMPLATE_PATH):
        return {"ok": False, "msg": f"No se encontró la plantilla en {TEMPLATE_PATH}", "pdf_path": None, "docx_path": None}


    # Carga el documento base DOCX
    doc = Document(TEMPLATE_PATH)


    # Datos del docente
    if docente and ("apellidos" in docente or "nombres" in docente):
        apellido = (docente.get("apellidos") or "").strip()
        nombre = (docente.get("nombres") or "").strip()
        docente_text = f"{apellido} {nombre}".strip()
    else:
        docente_text = ""


    # Prepara textos de reemplazo para la plantilla
    codigo_text = codigo_doc or ""
    grado_nombre, seccion_nombre = _grado_nombre_corto(grado_seccion)
    fecha_impresion = datetime.now().strftime("%d/%m/%Y")


    # Reemplazo de marcadores
    for marker, text in [("{DOCENTE}", docente_text), ("{CODIGO}", codigo_text),
                         ("{GRADO}", grado_nombre), ("{SECCION}", seccion_nombre),
                         ("{FECHA_IMPRESION}", fecha_impresion)]:
        _replace_text_in_paragraphs(doc, marker, text)
        _replace_in_tables(doc, marker, text)


    # Bandera para saber si se insertó la tabla
    tabla_insertada = False
    section = doc.sections[0]
    try:
        # Calcula ancho útil de la página
        page_width_in = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
    except Exception:
        page_width_in = 8.0


    # Insertar tabla de asistencias
    for p in doc.paragraphs:
        if "{TABLA_ASISTENCIAS}" in p.text:
            num_date_cols = len(fechas)
            num_cols = 3 + num_date_cols
            table = doc.add_table(rows=1, cols=num_cols)
            table.style = 'Table Grid'
            table.autofit = False
            table.allow_autofit = False
            table.alignment = WD_TABLE_ALIGNMENT.LEFT


            # Definir anchos
            nro_w = page_width_in * 0.07
            code_w = page_width_in * 0.16
            name_w = page_width_in * 0.47
            remaining = page_width_in - (nro_w + code_w + name_w)
            date_w = remaining / num_date_cols if num_date_cols > 0 else 0
            widths = [nro_w, code_w, name_w] + ([date_w] * num_date_cols)


            # Cabecera
            hdr_cells = table.rows[0].cells
            hdr_cells[0].paragraphs[0].add_run("Nro").bold = True
            hdr_cells[1].paragraphs[0].add_run("Código").bold = True
            hdr_cells[2].paragraphs[0].add_run("Nombre").bold = True
            for idx, f in enumerate(fechas):
                hdr_cells[3 + idx].paragraphs[0].add_run(f.strftime("%d/%m")).bold = True


            # Filas alumnos
            for nro, alumno in enumerate(alumnos, start=1):
                row_cells = table.add_row().cells
                # Nro
                row_cells[0].paragraphs[0].add_run(str(nro)).font.size = Pt(10)
                row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # Código
                row_cells[1].paragraphs[0].add_run(str(alumno["id_estudiante"])).font.size = Pt(10)
                row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # Nombre
                nombre_completo = f"{alumno.get('apellidos','')} {alumno.get('nombres','')}".strip()
                nombre_completo = " ".join(nombre_completo.split())
                row_cells[2].paragraphs[0].add_run(nombre_completo).font.size = Pt(10)
                row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                # Asistencias
                for j, fecha in enumerate(fechas):
                    estado = asist_map.get((alumno["id_matricula"], fecha), "")
                    if estado == "presente":
                        row_cells[3 + j].paragraphs[0].add_run("X").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


            # Limpia el marcador del párrafo donde estaba la tabla
            p.text = ""
            tabla_insertada = True
            break


    # Guardar DOCX
    try:
        # Intenta usar la carpeta Descargas del usuario
        downloads = os.path.join(os.path.expanduser("~"), "Downloads")
    except Exception:
        downloads = os.getcwd()


    # Construye nombre base del archivo de salida
    base_name = f"reporte_asistencias_{grado_seccion}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    docx_path = os.path.join(downloads, f"{base_name}.docx")
    doc.save(docx_path)


    # Convertir a PDF
    pdf_path = None
    try:
        from docx2pdf import convert
        pdf_path = os.path.join(downloads, f"{base_name}.pdf")
        convert(docx_path, pdf_path)
    except Exception:
        pdf_path = None


    # Mensaje final de resultado
    msg = "Generado correctamente."
    if pdf_path:
        return {"ok": True, "msg": msg, "pdf_path": pdf_path, "docx_path": docx_path}
    else:
        return {"ok": True, "msg": msg + " (PDF no generado automáticamente; se guardó DOCX)", "pdf_path": None, "docx_path": docx_path}
